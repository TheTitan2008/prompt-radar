"""Read-only DOCX extraction."""

from __future__ import annotations

import zipfile
from pathlib import Path

from prompt_radar.attachments.base import AttachmentExtraction, ExtractedSection
from prompt_radar.models import Attachment


class DocxAttachmentExtractor:
    """Extract paragraphs and tables without executing embedded content."""

    def extract(
        self, attachment: Attachment, path: Path
    ) -> AttachmentExtraction:
        if not zipfile.is_zipfile(path):
            return AttachmentExtraction(
                attachment_id=attachment.attachment_id,
                filename=attachment.filename,
                extraction_status="signature_mismatch",
                warnings=["DOCX is not a valid OOXML ZIP container"],
            )
        with zipfile.ZipFile(path) as archive:
            if "word/document.xml" not in set(archive.namelist()):
                return AttachmentExtraction(
                    attachment_id=attachment.attachment_id,
                    filename=attachment.filename,
                    extraction_status="signature_mismatch",
                    warnings=["OOXML Word document structure is missing"],
                )
        try:
            from docx import Document
        except ImportError:
            return AttachmentExtraction(
                attachment_id=attachment.attachment_id,
                filename=attachment.filename,
                extraction_status="dependency_missing",
                warnings=["Install the attachments dependency group"],
            )
        document = Document(path)
        text = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                text.append(" | ".join(cell.text for cell in row.cells))
        return AttachmentExtraction(
            attachment_id=attachment.attachment_id,
            filename=attachment.filename,
            extraction_status="success",
            sections=[
                ExtractedSection(
                    attachment_id=attachment.attachment_id,
                    filename=attachment.filename,
                    locator="document",
                    text="\n".join(text),
                )
            ],
        )

